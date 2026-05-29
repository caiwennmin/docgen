#!/usr/bin/env python3
"""办公自动化 - 模板 + 数据 → 批量生成 Word / Excel 文档"""

import io
import json
import os
import re
import shutil
import tempfile
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from zipfile import ZipFile

from flask import (
    Flask,
    jsonify,
    render_template,
    request,
    send_file,
    session,
)
from werkzeug.utils import secure_filename

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment

# ---------------------------------------------------------------------------
# config
# ---------------------------------------------------------------------------
BASE = Path(__file__).resolve().parent
UPLOAD_DIR = BASE / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "docgen-dev-key-change-in-production")
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB

# ---------------------------------------------------------------------------
# helpers – document generation
# ---------------------------------------------------------------------------

def _replace_placeholders_docx(doc: Document, row: dict) -> Document:
    """Replace {{key}} placeholders in a docx with values from a data row."""
    for para in doc.paragraphs:
        for run in para.runs:
            changed = False
            for key, val in row.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(val))
                    changed = True
            # handle combined runs
            if not changed:
                full_text = para.text
                replacements = False
                for key, val in row.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in full_text:
                        replacements = True
                if replacements:
                    new_text = full_text
                    for key, val in row.items():
                        new_text = new_text.replace("{{" + key + "}}", str(val))
                    # clear & set
                    for run in para.runs:
                        run.text = ""
                    para.runs[0].text = new_text if para.runs else ""

    # tables
    for table in doc.tables:
        for row_cells in table.rows:
            for cell in row_cells.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    for key, val in row.items():
                        full_text = full_text.replace("{{" + key + "}}", str(val))
                    if full_text != para.text:
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = full_text
    return doc


def generate_docx(template_path: Path, data_df: pd.DataFrame, output_dir: Path) -> list:
    """Generate one docx per row."""
    files = []
    for idx, (_, row) in enumerate(data_df.iterrows()):
        doc = Document(str(template_path))
        row_dict = {col: row[col] for col in data_df.columns if pd.notna(row[col])}
        doc = _replace_placeholders_docx(doc, row_dict)
        out_name = f"doc_{idx + 1}.docx"
        out_path = output_dir / out_name
        doc.save(str(out_path))
        files.append(out_path)
    return files


def generate_xlsx(template_path: Path, data_df: pd.DataFrame, output_dir: Path) -> list:
    """Fill an Excel template sheet-by-sheet using placeholder markers."""
    wb = load_workbook(str(template_path))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and "{{" in cell.value:
                    val = cell.value
                    for col, col_val in data_df.iloc[0].items():
                        val = val.replace("{{" + col + "}}", str(col_val))
                    cell.value = val

    # if more rows -> append
    if len(data_df) > 1:
        sheet = wb.active
        for row_idx, (_, data_row) in enumerate(data_df.iloc[1:].iterrows(), start=2):
            for col_idx, col_name in enumerate(data_df.columns, start=1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                cell.value = data_row[col_name]

    out_path = output_dir / "result.xlsx"
    wb.save(str(out_path))
    return [out_path]

# ---------------------------------------------------------------------------
# routes
# ---------------------------------------------------------------------------


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/upload-template", methods=["POST"])
def upload_template():
    file = request.files.get("file")
    if not file:
        return jsonify({"ok": False, "error": "no file"}), 400

    fname = secure_filename(file.filename)
    filepath = UPLOAD_DIR / f"tpl_{uuid.uuid4().hex[:8]}_{fname}"
    file.save(str(filepath))

    # detect placeholders
    placeholders = []
    ext = Path(fname).suffix.lower()
    if ext == ".docx":
        doc = Document(str(filepath))
        all_text = " ".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + " ".join(p.text for p in cell.paragraphs)
        placeholders = list(set(re.findall(r"\{\{(.+?)\}\}", all_text)))
    elif ext in (".xlsx", ".xls"):
        wb = load_workbook(str(filepath), data_only=True)
        all_text = ""
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        all_text += str(cell.value) + " "
        placeholders = list(set(re.findall(r"\{\{(.+?)\}\}", all_text)))

    return jsonify({
        "ok": True,
        "filename": fname,
        "filepath": str(filepath),
        "placeholders": placeholders,
    })


@app.route("/api/upload-data", methods=["POST"])
def upload_data():
    file = request.files.get("file")
    if not file:
        return jsonify({"ok": False, "error": "no file"}), 400

    fname = secure_filename(file.filename)
    filepath = UPLOAD_DIR / f"dat_{uuid.uuid4().hex[:8]}_{fname}"
    file.save(str(filepath))

    ext = Path(fname).suffix.lower()
    if ext in (".xlsx", ".xls"):
        df = pd.read_excel(str(filepath))
    elif ext == ".csv":
        df = pd.read_csv(str(filepath))
    else:
        return jsonify({"ok": False, "error": "unsupported data format"}), 400

    return jsonify({
        "ok": True,
        "filename": fname,
        "columns": df.columns.tolist(),
        "row_count": len(df),
        "filepath": str(filepath),
        "preview": df.head(5).fillna("").to_dict(orient="records"),
    })


@app.route("/api/preview", methods=["POST"])
def preview():
    """Generate one preview document (first row only)."""
    data = request.get_json()
    template_path = Path(data["template_path"])
    data_path = Path(data["data_path"])
    ext = template_path.suffix.lower()

    df = _read_data(data_path)
    preview_dir = UPLOAD_DIR / f"preview_{uuid.uuid4().hex[:8]}"
    preview_dir.mkdir()

    if ext == ".docx":
        files = generate_docx(template_path, df.head(1), preview_dir)
    elif ext in (".xlsx", ".xls"):
        files = generate_xlsx(template_path, df.head(1), preview_dir)
    else:
        return jsonify({"ok": False, "error": "unsupported"}), 400

    return jsonify({
        "ok": True,
        "file": str(files[0]),
        "filename": files[0].name,
    })


@app.route("/api/generate", methods=["POST"])
def generate():
    data = request.get_json()
    template_path = Path(data["template_path"])
    data_path = Path(data["data_path"])
    ext = template_path.suffix.lower()

    df = _read_data(data_path)
    task_id = uuid.uuid4().hex[:8]
    output_dir = UPLOAD_DIR / f"task_{task_id}"
    output_dir.mkdir()

    if ext == ".docx":
        files = generate_docx(template_path, df, output_dir)
    elif ext in (".xlsx", ".xls"):
        files = generate_xlsx(template_path, df, output_dir)
    else:
        return jsonify({"ok": False, "error": "unsupported template"}), 400

    # zip all
    zip_path = UPLOAD_DIR / f"task_{task_id}.zip"
    with ZipFile(str(zip_path), "w") as zf:
        for f in files:
            zf.write(str(f), f.name)

    # clean
    shutil.rmtree(str(output_dir), ignore_errors=True)

    return jsonify({
        "ok": True,
        "task_id": task_id,
        "file_count": len(files),
    })


@app.route("/api/download/<task_id>")
def download(task_id):
    zip_path = UPLOAD_DIR / f"task_{task_id}.zip"
    if not zip_path.exists():
        return jsonify({"ok": False, "error": "not found"}), 404
    return send_file(
        str(zip_path),
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"documents_{task_id}.zip",
    )


@app.route("/api/download-preview/<filename>")
def download_preview(filename):
    safe = secure_filename(filename)
    filepath = UPLOAD_DIR / safe
    if not filepath.exists():
        return jsonify({"ok": False, "error": "not found"}), 404
    return send_file(str(filepath), as_attachment=True, download_name=safe)


def _read_data(path: Path) -> pd.DataFrame:
    ext = path.suffix.lower()
    if ext in (".xlsx", ".xls"):
        return pd.read_excel(str(path))
    return pd.read_csv(str(path))


if __name__ == "__main__":
    app.run(debug=True, port=5050)


